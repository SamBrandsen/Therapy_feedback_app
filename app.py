import streamlit as st
from io import BytesIO
from docx import Document
from datetime import datetime, timezone
from collections import defaultdict

st.set_page_config(page_title="Therapy Reflection Template Generator", layout="centered")

# ------------------------------------------------
# QUESTION BANK
# ------------------------------------------------

QUESTION_BANK = [

# -------- Overall Experience --------
{
"id": "OE1",
"text": "On a scale of 1 to 5, how well is therapy meeting your needs? (Let 1 be 'not at all' and 5 be 'very well'.)",
"section": "Overall Experience",
"style": "structured",
"focus": "general"
},

{
"id": "OE7",
"text": "Overall, what emotions or thoughts come up when you think about therapy?",
"section": "Overall Experience",
"style": "unstructured",
"focus": "general"
},

{
"id": "OE2",
"text": "Is there anything about therapy that you especially appreciate or hope continues?",
"section": "Overall Experience",
"style": "both",
"focus": "general"
},

{
"id": "OE3",
"text": "Is there anything stressful, uncomfortable, or less effective about therapy?",
"section": "Overall Experience",
"style": "both",
"focus": "general"
},

{
"id": "OE5",
"text": "If you could change one thing about therapy, what would it be?",
"section": "Overall Experience",
"style": "structured",
"focus": "general"
},

{
"id": "OE6",
"text": "What do you experience your therapist as least aware of in how they engage with you?",
"section": "Overall Experience",
"style": "unstructured",
"focus": "general"
},

# -------- Topics & Comfort --------
{
"id": "TC1",
"text": "What topics do you find easy to discuss with your therapist?",
"section": "Topics & Comfort",
"style": "structured",
"focus": "general"
},

{
"id": "TC2",
"text": "What topics are more challenging to discuss? Are there any changes in approach that would help you feel more comfortable discussing them?",
"section": "Topics & Comfort",
"style": "structured",
"focus": "general"
},

{
"id": "TC3",
"text": "Are there any topics you would like your therapist to be more proactive in discussing with you, or topics you have been hesitant to bring up?",
"section": "Topics & Comfort",
"style": "both",
"focus": "general"
},

{
"id": "TC4",
"text": "Overall, are there topics that tend to be more or less comfortable for you to discuss with your therapist?",
"section": "Topics & Comfort",
"style": "unstructured",
"focus": "general"
},

# -------- Structure & Style --------
{
"id": "SS1",
"text": "Is there anything you would like to discuss about punctuality in sessions or how sessions are scheduled?",
"section": "Structure & Style",
"style": "structured",
"focus": "general"
},

{
"id": "SS4",
"text": "How do you feel about the amount or type of self-disclosure in sessions (how much your therapist shares about themselves)?",
"section": "Structure & Style",
"style": "structured",
"focus": "general"
},

{
"id": "SS5",
"text": "How do you feel about the amount of structure or direction from the therapist during sessions?",
"section": "Structure & Style",
"style": "structured",
"focus": "general"
},

{
"id": "SS2",
"text": "How do you experience the structure of sessions (punctuality, disclosure, how sessions flow)?",
"section": "Structure & Style",
"style": "unstructured",
"focus": "general"
},

{
"id": "SS3",
"text": "Is there anything you would like to change about the physical environment of therapy (lighting, seating, proximity to door)?",
"section": "Structure & Style",
"style": "both",
"focus": "general"
},

{
"id": "SS6",
"text": "Is there anything else about the structure or setting of therapy that you would like to mention?",
"section": "Structure & Style",
"style": "both",
"focus": "general"
},

# -------- Relational Climate --------
{
"id": "RC1",
"text": "How safe, respected, and understood do you feel in therapy? Is there anything that would increase your sense of safety?",
"section": "Relational Climate",
"style": "structured",
"focus": "general"
},

{
"id": "RC2",
"text": "How comfortable do you feel bringing up feedback or difficult topics with your therapist? Is there anything that would help you feel more comfortable?",
"section": "Relational Climate",
"style": "both",
"focus": "general"
},

{
"id": "RC3",
"text": "What emotional experiences do you tend to have during or after sessions?",
"section": "Relational Climate",
"style": "unstructured",
"focus": "general"
},

{
"id": "RC4",
"text": "Are there any emotional patterns around therapy that you would like your therapist to know about (for example anxiety before sessions or relief afterward)?",
"section": "Relational Climate",
"style": "structured",
"focus": "general"
},

# -------- Therapy Harm & Boundaries --------
{
"id": "TH1",
"text": "Are there moments where you feel confused about your role in therapy (for example feeling treated more like a friend than a client, or other role confusion)?",
"section": "Therapy Harm & Boundaries",
"style": "structured",
"focus": "harm"
},

{
"id": "TH16",
"text": "Has contact (or lack of contact) between sessions ever been a source of confusion or pain (texts, emails, calls)? What changes would help you feel most safe?",
"section": "Therapy Harm & Boundaries",
"style": "structured",
"focus": "harm"
},

{
"id": "TH2",
"text": "Are there moments where you feel dismissed or unheard in therapy?",
"section": "Therapy Harm & Boundaries",
"style": "both",
"focus": "harm"
},

{
"id": "TH9",
"text": "Are there parts of therapy that feel especially confusing or painful (for example inconsistent boundaries, being overly warm or withdrawn)?",
"section": "Therapy Harm & Boundaries",
"style": "both",
"focus": "harm"
},

{
"id": "TH10",
"text": "Are there topics or aspects of your identity where your therapist's responses have felt hurtful?",
"section": "Therapy Harm & Boundaries",
"style": "both",
"focus": "harm"
},

{
"id": "TH11",
"text": "Is there anything you would like your therapist to understand about how therapy harm has affected you?",
"section": "Therapy Harm & Boundaries",
"style": "both",
"focus": "harm"
},

{
"id": "TH21",
"text": "What response from your therapist would feel most healing for you?",
"section": "Therapy Harm & Boundaries",
"style": "both",
"focus": "harm"
},

{
"id": "TH22",
"text": "Are there parts of your feedback that you worry your therapist might misunderstand?",
"section": "Therapy Harm & Boundaries",
"style": "both",
"focus": "harm"
},

]

# ------------------------------------------------
# FINAL QUESTION
# ------------------------------------------------

FINAL_QUESTION = {
"text": "How would you like your therapist to use this feedback? (For example: discussing it together in session, reading it privately, taking time to reflect before responding, or exploring other support such as supervision or third party facilitation.)"
}

# ------------------------------------------------
# FILTERING
# ------------------------------------------------

def filter_questions(style, focus):
    style = style.lower()  # normalize
    selected = []

    for q in QUESTION_BANK:

        if q["style"] not in [style, "both"]:
            continue

        if focus == "General Feedback" and q["focus"] != "general":
            continue
        elif focus == "Concerns / Harm" and q["focus"] != "harm":
            continue
        elif focus == "Both":
            pass  # include all

        selected.append(q)

    return selected

# ------------------------------------------------
# GROUPING
# ------------------------------------------------

SECTION_ORDER = [
    "Overall Experience",
    "Topics & Comfort",
    "Structure & Style",
    "Relational Climate",
    "Therapy Harm & Boundaries"
]

def group_by_section(questions):

    sections = defaultdict(list)

    for q in questions:
        sections[q["section"]].append(q["text"])

    ordered_sections = {}

    for section in SECTION_ORDER:
        if section in sections:
            ordered_sections[section] = sections[section]

    return ordered_sections

# ------------------------------------------------
# DOCUMENT GENERATION
# ------------------------------------------------

def make_docx(sections):

    doc = Document()

    doc.add_heading("Therapy Reflection Template", level=1)

    doc.add_paragraph(
        f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}"
    )

    doc.add_paragraph(
        "This document is generated through a tool that aims to help clients share their experiences in therapy with their therapist. "
        "The app includes a broad range of question prompts and we hope that this tool may provide a helpful starting point for clinicians and clients "
        "in checking in about the therapeutic process."
    )

    for section, questions in sections.items():

        doc.add_heading(section, level=2)

        for q in questions:

            doc.add_paragraph(q)

            for _ in range(3):
                doc.add_paragraph("________________________________________________________")

            doc.add_paragraph()

    doc.add_heading("Sharing Feedback", level=2)

    doc.add_paragraph(FINAL_QUESTION["text"])

    for _ in range(4):
        doc.add_paragraph("________________________________________________________")

    f = BytesIO()
    doc.save(f)
    f.seek(0)

    return f.read()

# ------------------------------------------------
# UI
# ------------------------------------------------

st.title("Therapy Reflection Template Generator")

st.markdown("""
### About the Therapy Reflection Template Generator

This tool aims to help you reflect on your experience in therapy — what feels helpful, what feels difficult, and what you might want your therapist to understand.

The tool generates a blank reflection template based on the preferences you select. After downloading the document, you may answer as many or as few questions as you like.

This website does not collect or store responses. It only generates a downloadable template for your own use.

---

### How You Might Use this Tool

People use reflection tools like this in different ways. Some use them to organize their own thoughts. Others choose to share some or all of their reflections with a therapist or another trusted person.

Our hope is that this tool can assist any client who feels like they may benefit from sharing feedback with their therapist, but who may need some support in order to do so.

Conversations about feedback can sometimes strengthen a therapy relationship. At the same time, you are never required to share feedback if doing so feels uncomfortable or unsafe.

---

### Disclaimer

Please also note that materials shared with a therapist may become part of the clinical record depending on their documentation practices.

Please also be aware that in circumstances of therapy harm or abuse, it is not always safe for a client to give feedback to a therapist. If you are experiencing therapy harm or abuse, or wonder if you may be, **click here for more information and resources**.

---

### About Us

This tool was created by **Sam Brandsen** in collaboration with the **Therapy Harm Response & Prevention initiative**.

If you have any feedback about the app, please feel welcome to contact us at **sambrandsen7@gmail.com**.
""")

focus = st.radio(
    "Focus of reflection",
    ["General Feedback", "Concerns / Harm", "Both"]
)

style = st.radio(
    "Question style",
    ["Structured", "Unstructured"]
)

with st.expander("What is the difference between structured and unstructured questions?"):

    st.markdown("""
**Structured questions**
- Focus on specific aspects of therapy  
- Often easier to answer quickly  

**Unstructured questions**
- Broader reflection prompts  
- Allow more open description of your experience
""")

questions = filter_questions(style, focus)
sections = group_by_section(questions)

if st.button("Generate Reflection Template"):

    if not questions:
        st.warning("No questions match the selected settings.")
    else:

        doc = make_docx(sections)

        filename = f"therapy_reflection_template_{datetime.now().strftime('%Y%m%d')}.docx"

        st.download_button(
            "Download Word Template",
            doc,
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
